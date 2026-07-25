from docx import Document

from jobagent.ai.tailor import ExperienceItem, SkillGroup, TailoredResume
from jobagent.docgen import slugify, write_cover_letter_docx, write_resume_docx


def sample_resume() -> TailoredResume:
    return TailoredResume(
        name="Chris Leveque",
        contact="City, ST · chris@example.com · (555) 555-5555",
        summary="",
        skills=["Okta", "Entra ID", "SailPoint"],
        skill_groups=[
            SkillGroup(name="Cybersecurity", items=["Okta", "Entra ID"]),
            SkillGroup(name="Technical Support", items=["SailPoint"]),
        ],
        experience=[
            ExperienceItem(company="Acme", title="IAM Engineer",
                           dates="2021 – Present", location="Remote",
                           bullets=["Migrated 5k users to Okta", "Automated JML with SCIM"]),
        ],
        projects=["Built a home SOC lab with Wazuh"],
        education=["B.S. Something, Some University"],
        certifications=["Okta Certified Professional"],
    )


def sample_contact() -> dict:
    return {
        "full_name": "Chris Leveque", "city": "Austin, TX",
        "phone": "(555) 555-5555", "email": "chris@example.com",
        "linkedin": "https://linkedin.com/in/chris-l",
        "github": "https://github.com/chrisl",
        "tryhackme": "https://tryhackme.com/p/chrisl",
    }


def all_text(path) -> str:
    """Collect text from the running header, body paragraphs, and table cells."""
    doc = Document(str(path))
    parts = [p.text for p in doc.sections[0].header.paragraphs]
    parts += [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts += [p.text for p in cell.paragraphs]
    return "\n".join(parts)


def hyperlink_targets(path) -> list[str]:
    """Hyperlink relationship targets across the document and header parts."""
    doc = Document(str(path))
    targets = []
    for part in (doc.part, doc.sections[0].header.part):
        targets += [rel.target_ref for rel in part.rels.values()
                    if rel.reltype.endswith("/hyperlink")]
    return targets


def test_write_resume_docx(tmp_path):
    path = write_resume_docx(sample_resume(), tmp_path / "out" / "resume.docx",
                             contact=sample_contact())
    assert path.exists()
    text = all_text(path)
    for expected in ("Chris Leveque", "EXPERIENCE", "Acme", "IAM Engineer",
                     "Migrated 5k users to Okta", "KEY COMPETENCIES",
                     "Cybersecurity", "TECHNICAL PROJECTS",
                     "Built a home SOC lab with Wazuh",
                     "Okta Certified Professional",
                     "B.S. Something, Some University"):
        assert expected in text


def test_resume_uses_two_column_table(tmp_path):
    path = write_resume_docx(sample_resume(), tmp_path / "resume.docx",
                             contact=sample_contact())
    doc = Document(str(path))
    assert len(doc.tables) == 1
    row = doc.tables[0].rows[0]
    assert len(row.cells) == 2
    left = "\n".join(p.text for p in row.cells[0].paragraphs)
    right = "\n".join(p.text for p in row.cells[1].paragraphs)
    assert "EXPERIENCE" in left and "TECHNICAL PROJECTS" in left
    assert "KEY COMPETENCIES" in right and "CERTIFICATIONS" in right
    assert "EDUCATION" in right


def test_header_has_labeled_hyperlinks(tmp_path):
    path = write_resume_docx(sample_resume(), tmp_path / "resume.docx",
                             contact=sample_contact())
    targets = hyperlink_targets(path)
    assert "mailto:chris@example.com" in targets
    assert "https://linkedin.com/in/chris-l" in targets
    assert "https://github.com/chrisl" in targets
    assert "https://tryhackme.com/p/chrisl" in targets
    # plain segments must NOT become links
    assert not any("Austin" in t or "555" in t for t in targets)
    # links render as clean labels in the header, not raw URLs
    header_text = "\n".join(
        p.text for p in Document(str(path)).sections[0].header.paragraphs)
    for label in ("Github", "LinkedIn", "TryHackMe"):
        assert label in header_text
    assert "linkedin.com/in/chris-l" not in header_text
    assert "(555) 555-5555" in header_text


def test_contact_fallback_when_no_contact_dict(tmp_path):
    resume = sample_resume()
    resume.contact = ("Austin, TX · (555) 555-5555 · chris@example.com · "
                      "linkedin.com/in/chris-l · github.com/chrisl")
    path = write_resume_docx(resume, tmp_path / "resume.docx")
    targets = hyperlink_targets(path)
    assert "mailto:chris@example.com" in targets
    assert "https://linkedin.com/in/chris-l" in targets
    text = all_text(path)
    assert "Austin, TX" in text and "(555) 555-5555" in text


def test_flat_skills_used_without_groups(tmp_path):
    resume = sample_resume()
    resume.skill_groups = []
    path = write_resume_docx(resume, tmp_path / "resume.docx")
    text = all_text(path)
    for skill in ("Okta", "Entra ID", "SailPoint"):
        assert skill in text


def test_write_cover_letter_docx(tmp_path):
    letter = "Dear Hiring Manager,\n\nFirst paragraph.\n\nSecond paragraph."
    path = write_cover_letter_docx(letter, "Chris Leveque", tmp_path / "cl.docx",
                                   contact=sample_contact())
    text = all_text(path)
    assert "Dear Hiring Manager," in text
    assert "Second paragraph." in text
    assert "Chris Leveque" in text


def test_slugify():
    assert slugify("Sr. IAM Engineer @ Acme, Inc.") == "sr-iam-engineer-acme-inc"
    assert slugify("///") == "job"


def test_compact_levels_tighten_layout(tmp_path):
    from docx.shared import Inches, Pt

    loose = Document(str(write_resume_docx(
        sample_resume(), tmp_path / "loose.docx", contact=sample_contact())))
    tight = Document(str(write_resume_docx(
        sample_resume(), tmp_path / "tight.docx", contact=sample_contact(),
        compact=2)))
    assert tight.styles["Normal"].font.size == Pt(9.5)
    assert loose.styles["Normal"].font.size == Pt(10)
    assert tight.sections[0].left_margin == Inches(0.6)
    assert loose.sections[0].left_margin == Inches(0.7)


def test_table_borders_declared_none_before_layout(tmp_path):
    import re as _re
    import zipfile

    path = write_resume_docx(sample_resume(), tmp_path / "resume.docx",
                             contact=sample_contact())
    xml = zipfile.ZipFile(str(path)).read("word/document.xml").decode()
    tblpr = _re.search(r"<w:tblPr>.*?</w:tblPr>", xml, _re.S).group(0)
    # every border explicitly none, positioned before tblLayout so strict
    # parsers (Windows LibreOffice) honor it
    assert tblpr.count('w:val="none"') == 6
    assert tblpr.index("tblBorders") < tblpr.index("tblLayout")
