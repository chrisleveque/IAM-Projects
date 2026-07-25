"""Generate .docx (and optionally .pdf via LibreOffice) from a tailored package.

Layout mirrors the user's original resume: a centered navy name header with
phone / email / labeled links and a rule (repeated on every page via the
running header), then a borderless two-column table — experience and projects
on the left, competencies / certifications / education on the right.
"""

from __future__ import annotations

import re
import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

NAVY = RGBColor(0x1F, 0x38, 0x64)
LINK_BLUE = "0563C1"
BODY_FONT = "Times New Roman"
BODY_SIZE = Pt(10)

# Matches URL-ish contact segments like linkedin.com/in/me, github.com/me,
# tryhackme.com/p/me, https://example.io — but not phones or "City, ST".
_URL_RE = re.compile(r"^(https?://)?(www\.)?[\w-]+(\.[\w-]+)+(/\S*)?$", re.I)

# Known sites render as a clean labeled link ("LinkedIn") instead of the raw URL.
_LINK_LABELS = (
    ("linkedin.com", "LinkedIn"),
    ("github.com", "Github"),
    ("tryhackme.com", "TryHackMe"),
    ("hackthebox.com", "HackTheBox"),
    ("credly.com", "Credly"),
)

# Header link order matches the original resume: Github | LinkedIn | TryHackMe
_HEADER_LINK_KEYS = (
    ("github", "Github"),
    ("linkedin", "LinkedIn"),
    ("tryhackme", "TryHackMe"),
    ("hackthebox", "HackTheBox"),
    ("credly", "Credly"),
)


def slugify(text: str, max_len: int = 40) -> str:
    slug = re.sub(r"[^a-z0-9]+", "-", text.lower()).strip("-")
    return slug[:max_len] or "job"


def _link_label(url: str) -> str:
    low = url.lower()
    for domain, label in _LINK_LABELS:
        if domain in low:
            return label
    return url


def _base_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE
    # Word's template defaults (1.08 line spacing, 8pt after each paragraph,
    # 1" top/bottom margins) balloon the page count — bring them down to the
    # original resume's density, but keep a little air between lines.
    style.paragraph_format.line_spacing = 1.05
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    section = doc.sections[0]
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.5)
    section.header_distance = Inches(0.3)


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    """python-docx has no high-level hyperlink API; build the XML directly."""
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), LINK_BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.append(color)
    props.append(underline)
    run.append(props)
    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _bottom_rule(paragraph) -> None:
    """Give a paragraph a thin bottom border (the header divider line)."""
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "444444")
    borders.append(bottom)
    p_pr.append(borders)


def _center(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(2)
    return paragraph


def _write_contact_segments(paragraph, contact: str) -> None:
    """Render a '·'/'|'-separated contact string with labeled hyperlinks."""
    first = True
    for segment in (s.strip() for s in re.split(r"\s*[·|]\s*", contact)):
        if not segment:
            continue
        if not first:
            paragraph.add_run("  ·  ")
        first = False
        if "@" in segment and " " not in segment and not segment.lower().startswith(
                ("http", "www")):
            _add_hyperlink(paragraph, segment, f"mailto:{segment}")
        elif _URL_RE.match(segment):
            url = segment if segment.lower().startswith("http") else f"https://{segment}"
            _add_hyperlink(paragraph, _link_label(segment), url)
        else:
            paragraph.add_run(segment)


def _build_running_header(doc: Document, name: str, contact: dict,
                          contact_fallback: str = "") -> None:
    """Centered name / phone / email / links block with a rule, repeated on
    every page (it lives in the document's running header, like the original)."""
    header = doc.sections[0].header
    name_par = _center(header.paragraphs[0])
    name_par.text = ""
    run = name_par.add_run(name)
    run.bold = True
    run.font.name = BODY_FONT
    run.font.size = Pt(24)
    run.font.color.rgb = NAVY

    if contact:
        phone = contact.get("phone")
        if phone:
            _center(header.add_paragraph()).add_run(str(phone)).bold = True
        email = contact.get("email")
        if email:
            _add_hyperlink(_center(header.add_paragraph()), str(email),
                           f"mailto:{email}")
        links = [(label, contact.get(key))
                 for key, label in _HEADER_LINK_KEYS if contact.get(key)]
        if links:
            links_par = _center(header.add_paragraph())
            for i, (label, url) in enumerate(links):
                if i:
                    links_par.add_run(" | ")
                _add_hyperlink(links_par, label, str(url))
    elif contact_fallback:
        _write_contact_segments(_center(header.add_paragraph()), contact_fallback)

    _bottom_rule(header.add_paragraph())


def _clear_table_borders(table) -> None:
    """Explicitly mark every table border as none so no viewer draws boxes."""
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        borders.append(el)
    table._tbl.tblPr.append(borders)


def _cell_writer(cell):
    """Cells start with one empty paragraph — reuse it for the first write."""
    state = {"used_first": False}

    def new_paragraph():
        if not state["used_first"]:
            state["used_first"] = True
            return cell.paragraphs[0]
        return cell.add_paragraph()

    return new_paragraph


def _section_heading(new_par, text: str) -> None:
    par = new_par()
    par.paragraph_format.space_before = Pt(8)
    par.paragraph_format.space_after = Pt(4)
    run = par.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = NAVY


def _subheading(new_par, text: str) -> None:
    par = new_par()
    par.paragraph_format.space_before = Pt(6)
    par.paragraph_format.space_after = Pt(2)
    run = par.add_run(text)
    run.bold = True
    run.font.color.rgb = NAVY


def _bullet(cell, text: str) -> None:
    # Manual hanging-indent bullets: the built-in List Bullet style indents
    # ~0.5" before text starts, which wraps nearly every line an extra time
    # inside a table column and pushes the resume past two pages.
    par = cell.add_paragraph()
    par.paragraph_format.left_indent = Inches(0.14)
    par.paragraph_format.first_line_indent = Inches(-0.14)
    par.paragraph_format.space_after = Pt(2)
    par.add_run("• " + text)


def _left_column(cell, resume) -> None:
    new_par = _cell_writer(cell)
    _section_heading(new_par, "Experience")
    for role in resume.experience:
        company_par = new_par()
        company_par.paragraph_format.space_before = Pt(6)
        company_par.paragraph_format.space_after = Pt(0)
        company_par.add_run(role.company).bold = True
        if role.location:
            company_par.add_run(f" – {role.location}")
        title_par = new_par()
        title_par.paragraph_format.space_after = Pt(2)
        title_line = role.title + (f" | {role.dates}" if role.dates else "")
        title_par.add_run(title_line)
        for bullet in role.bullets:
            _bullet(cell, bullet)
    if resume.projects:
        _section_heading(new_par, "Technical Projects")
        for project in resume.projects:
            _bullet(cell, project)


def _right_column(cell, resume) -> None:
    new_par = _cell_writer(cell)
    _section_heading(new_par, "Key Competencies")
    if resume.skill_groups:
        for group in resume.skill_groups:
            _subheading(new_par, group.name)
            for item in group.items:
                _bullet(cell, item)
    else:
        for skill in resume.skills:
            _bullet(cell, skill)
    if resume.certifications:
        _section_heading(new_par, "Certifications")
        for cert in resume.certifications:
            _bullet(cell, cert)
    if resume.education:
        _section_heading(new_par, "Education")
        for line in resume.education:
            par = new_par()
            par.paragraph_format.space_after = Pt(2)
            par.add_run(line)


def write_resume_docx(resume, path: Path, contact: dict | None = None) -> Path:
    doc = Document()
    _base_style(doc)
    _build_running_header(doc, resume.name, contact or {},
                          contact_fallback=resume.contact)

    if resume.summary:
        summary = doc.add_paragraph(resume.summary)
        summary.paragraph_format.space_after = Pt(6)

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    _clear_table_borders(table)
    left, right = table.rows[0].cells
    left.width = Inches(4.85)
    right.width = Inches(2.25)
    _left_column(left, resume)
    _right_column(right, resume)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path


def write_cover_letter_docx(text: str, name: str, path: Path,
                            contact: dict | None = None) -> Path:
    doc = Document()
    _base_style(doc)
    _build_running_header(doc, name, contact or {})
    for para_text in (p.strip() for p in text.split("\n\n") if p.strip()):
        par = doc.add_paragraph(para_text)
        par.paragraph_format.space_after = Pt(8)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path


def find_soffice() -> str | None:
    """Locate LibreOffice even when it isn't on PATH (its Windows installer
    doesn't add itself there)."""
    found = shutil.which("soffice") or shutil.which("libreoffice")
    if found:
        return found
    for candidate in (
        Path(r"C:\Program Files\LibreOffice\program\soffice.exe"),
        Path(r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"),
        Path("/Applications/LibreOffice.app/Contents/MacOS/soffice"),
    ):
        if candidate.exists():
            return str(candidate)
    return None


def convert_to_pdf(docx_path: Path) -> Path | None:
    """Convert with LibreOffice if installed; return the PDF path or None."""
    soffice = find_soffice()
    if soffice is None:
        return None
    pdf = docx_path.with_suffix(".pdf")
    # Delete any PDF from a previous run first — otherwise a failed conversion
    # would leave the stale file behind and we'd report it as fresh output.
    pdf.unlink(missing_ok=True)
    # A dedicated user profile lets headless conversion run even while the
    # user has LibreOffice open (it normally refuses: one instance per profile).
    profile = Path(tempfile.gettempdir()) / "jobagent-soffice-profile"
    try:
        subprocess.run(
            [soffice, "--headless",
             f"-env:UserInstallation={profile.absolute().as_uri()}",
             "--convert-to", "pdf", "--outdir", str(docx_path.parent),
             str(docx_path)],
            check=True, capture_output=True, timeout=120,
        )
    except (subprocess.SubprocessError, OSError):
        return None
    return pdf if pdf.exists() else None
